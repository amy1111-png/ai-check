import streamlit as st
import requests
import json
import pdfplumber
import pandas as pd
from docx import Document
from PIL import Image
import io
import base64
import time

# 網頁基本設定
st.set_page_config(page_title="AI 財務稽核與 Word 導出", layout="wide")
st.title("⚖️ AI 財務全能稽核系統 (全自動偵測版)")

# --- 側邊欄：API 與說明 ---
with st.sidebar:
    st.header("🔑 系統設定")
    api_key = st.text_input("請貼上 API Key", type="password").strip()
    st.divider()
    st.info("💡 建議同時上傳：\n1. 簽呈檔案\n2. 所有的廠商報價單(PDF或照片)\n3. 114年舊資料")

# --- 檔案上傳區 ---
uploaded_files = st.file_uploader(
    "上傳所有相關檔案", 
    type=['pdf', 'docx', 'xlsx', 'xls', 'png', 'jpg', 'jpeg'], 
    accept_multiple_files=True
)

def process_file(f):
    fname = f.name.lower()
    try:
        if fname.endswith('.pdf'):
            with pdfplumber.open(f) as pdf:
                return f"\n--- [文件內容: {f.name}] ---\n" + "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()]), None
        elif fname.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(f)
            return f"\n--- [表格內容: {f.name}] ---\n{df.to_string()}", None
        elif fname.endswith('.docx'):
            doc = Document(f)
            return f"\n--- [Word 內容: {f.name}] ---\n" + "\n".join([p.text for p in doc.paragraphs]), None
        elif fname.endswith(('.png', '.jpg', '.jpeg')):
            img = Image.open(f).convert('RGB')
            img.thumbnail((1200, 1200)) 
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=75)
            return f"\n[附件圖片: {f.name}]\n", base64.b64encode(buf.getvalue()).decode()
    except Exception as e:
        return f"讀取 {f.name} 失敗: {str(e)}", None
    return "", None

# --- 核心功能：產出 Word 檔案 ---
def create_word_report(content):
    doc = Document()
    doc.add_heading('AI 財務稽核報告', 0)
    doc.add_paragraph(f"產出日期: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("-" * 30)
    
    for line in content.split('\n'):
        if line.startswith('###'):
            doc.add_heading(line.replace('###', '').strip(), level=2)
        elif line.startswith('##'):
            doc.add_heading(line.replace('##', '').strip(), level=1)
        else:
            doc.add_paragraph(line)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 主分析邏輯 ---
if uploaded_files:
    all_texts = []
    all_imgs = []
    
    with st.spinner("🔍 正在讀取所有檔案與報價單..."):
        for f in uploaded_files:
            txt, img_b64 = process_file(f)
            if txt: all_texts.append(txt)
            if img_b64: all_imgs.append(img_b64)

    full_context = "\n\n".join(all_texts)
    
    tab1, tab2 = st.tabs(["📝 檔案數據預覽", "📊 AI 稽核勾稽報告"])
    
    with tab1:
        st.subheader("📁 已提取數據匯總")
        st.text_area("合併內容：", full_context, height=300)

    with tab2:
        if st.button("🚀 開始深度勾稽與年度分析", type="primary"):
            if not api_key:
                st.error("請輸入 API Key")
            else:
                with st.status("🛸 AI 正在進行數據驗證...", expanded=True) as status:
                    # 修正：直接鎖定正確的模型路徑
                    final_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
                    
                    # 修正：把寫死的 $489,000 拿掉，改為通用稽核邏輯
                    prompt = f"""
                    你是一位專業資深財務稽核。請針對提供的多份文件進行「三方勾稽」與「跨年度比對」。

                    【任務重點】：
                    1. **自動識別總額**：從簽呈檔案中識別本年度的預算總額，並核對與報價單上的金額是否完全相符。
                    2. **計算驗證**：自動加總報價單中的細項金額，檢查是否等於簽呈上的總額。
                    3. **年度對比**：找出 114 年與 115 年的金額差異，計算增減額與百分比。
                    4. **異常偵測**：若發現報價單與簽呈數字不符、單價異常、或 OCR 辨識錯字，請用【‼️錯誤警告】標示。
                    5. **費用分攤**：分析費用在公司與診所之間的攤提比例。

                    【報告格式】：
                    - 財務審核結論 (包含年度增減對比)
                    - 勾稽驗證細節 (廠商報價單 vs 簽呈)
                    - 文字與計算糾錯 (例如品項名稱筆誤)
                    - 建議詢問台詞 (針對疑點)

                    以下為檔案內容：
                    {full_context}
                    """
                    
                    user_parts = [{"text": prompt}]
                    for b64 in all_imgs:
                        user_parts.append({"inline_data": {"mime_type": "image/jpeg", "data": b64}})
                    
                    try:
                        res = requests.post(final_url, json={"contents": [{"parts": user_parts}]}, timeout=120)
                        
                        if res.status_code == 200:
                            status.update(label="✅ 分析完成！", state="complete")
                            result_text = res.json()['candidates'][0]['content']['parts'][0]['text']
                            st.markdown(result_text)
                            
                            st.divider()
                            st.subheader("📥 匯出專業報告")
                            word_data = create_word_report(result_text)
                            
                            st.download_button(
                                label="📥 下載 Word 稽核報告 (.docx)",
                                data=word_data,
                                file_name=f"財務稽核報告_{time.strftime('%Y%m%d')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        else:
                            st.error(f"分析失敗：{res.text}")
                    except Exception as e:
                        st.error(f"連線異常：{str(e)}")
